"""
Professional DOCX Generator for ADD_Scarecrow_Drone.md

Creates a professionally formatted Word document matching the style of:
- Professional cover page
- Table of contents with clickable links
- Clean tables with borders
- ASCII art diagrams preserved in monospace
- Consistent spacing and heading hierarchy
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def create_cover_page(doc):
    """Create a professional cover page."""
    # Add blank paragraphs for vertical centering
    for _ in range(8):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Application Design Document (ADD)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scarecrow Drone - Bird Detection System")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    doc.add_paragraph()

    # Version
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version: ")
    run.bold = True
    run.font.size = Pt(12)
    run = version.add_run("1.0")
    run.font.size = Pt(12)

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Date: ")
    run.bold = True
    run.font.size = Pt(12)
    run = date_p.add_run("January 2026")
    run.font.size = Pt(12)

    # Project
    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run("Project: ")
    run.bold = True
    run.font.size = Pt(12)
    run = project.add_run("Scarecrow Drone - Secondary Alpha Implementation")
    run.font.size = Pt(12)

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph for internal linking."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r

    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.insert(0, start)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)

def add_hyperlink(paragraph, text, bookmark_name):
    """Add a hyperlink to an internal bookmark."""
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create run inside hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rPr.append(rFonts)

    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
    rPr.append(sz)

    new_run.append(rPr)

    # Text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_toc(doc):
    """Create table of contents page with clickable links."""
    add_page_break(doc)

    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    doc.add_paragraph()

    toc_items = [
        ("1. Use Cases", "section_1"),
        ("2. System Architecture", "section_2"),
        ("3. Data Model", "section_3"),
        ("4. Behavioral Analysis", "section_4"),
        ("5. Object-Oriented Model", "section_5"),
        ("6. User Interface Draft", "section_6"),
        ("7. Testing", "section_7")
    ]

    for item_text, bookmark_name in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_hyperlink(p, item_text, bookmark_name)

def setup_styles(doc):
    """Set up custom styles for the document."""
    styles = doc.styles

    # Heading 1 style
    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 style
    h2 = styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(10)

    # Heading 3 style
    h3 = styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(8)

    # Heading 4 style
    h4 = styles['Heading 4']
    h4.font.name = 'Arial'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)

    # Normal style
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    header_row = table.rows[0]
    set_row_cant_split(header_row)  # Prevent header from splitting across pages
    header_cells = header_row.cells
    for idx, header in enumerate(headers):
        if idx < len(header_cells):
            header_text = header.strip()
            # Remove markdown bold syntax
            header_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', header_text)
            header_cells[idx].text = header_text
            # Bold header
            for paragraph in header_cells[idx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
            # Light gray background for header
            set_cell_shading(header_cells[idx], 'E0E0E0')

    # Data rows
    for row_data in rows:
        new_row = table.add_row()
        set_row_cant_split(new_row)  # Prevent row from splitting across pages
        row_cells = new_row.cells
        for idx, cell_text in enumerate(row_data):
            if idx < len(row_cells):
                # Handle multi-line cell content (br tags)
                cell_text = cell_text.strip()
                cell_text = cell_text.replace('<br>', '\n')
                cell_text = cell_text.replace('&nbsp;', ' ')
                # Remove markdown bold syntax
                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                row_cells[idx].text = cell_text
                for paragraph in row_cells[idx].paragraphs:
                    paragraph.paragraph_format.space_before = Pt(4)
                    paragraph.paragraph_format.space_after = Pt(4)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'

    doc.add_paragraph()  # Space after table

def add_code_block(doc, code_text):
    """Add a code block (for ASCII diagrams) to the document."""
    # Split by lines to handle long code blocks better
    lines = code_text.split('\n')

    # Determine if it's a diagram or OCL code
    is_ocl = 'context' in code_text.lower() or 'pre:' in code_text.lower() or 'post:' in code_text.lower() or 'inv:' in code_text.lower()

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8) if not is_ocl else Pt(9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

def parse_markdown(md_content):
    """Parse markdown content and convert to structured data."""
    lines = md_content.split('\n')
    elements = []

    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_headers = []
    table_rows = []
    skip_google_docs_format = False
    skip_markdown_toc = False

    while i < len(lines):
        line = lines[i]

        # Skip the markdown Table of Contents section (we create our own)
        if line.strip() == '## Table of Contents':
            skip_markdown_toc = True
            i += 1
            continue

        if skip_markdown_toc:
            # End of TOC when we hit the next major section or ---
            if line.startswith('## ') and 'Table of Contents' not in line:
                skip_markdown_toc = False
            elif line.strip() == '---':
                skip_markdown_toc = False
                i += 1
                continue
            else:
                i += 1
                continue

        # Skip Google Docs format sections (alternative ASCII diagrams)
        if '**Google Docs Format:**' in line or 'Google Docs Format:' in line:
            skip_google_docs_format = True
            i += 1
            continue

        # Check if we're at a section header or '---' to end Google Docs skip
        if skip_google_docs_format:
            if line.startswith('##') or line.startswith('###') or line.strip() == '---':
                skip_google_docs_format = False
            else:
                i += 1
                continue

        # Handle code blocks (for ASCII diagrams)
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and not line.strip().startswith('#'):
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells from split

            if not in_table:
                in_table = True
                table_headers = cells
            elif re.match(r'^[\s|:\-]+$', line.replace('|', '').strip()):
                # Skip separator line
                pass
            else:
                table_rows.append(cells)

            i += 1
            continue
        elif in_table:
            # End of table
            if table_headers:
                elements.append(('table', (table_headers, table_rows)))
            table_headers = []
            table_rows = []
            in_table = False

        # Handle headings
        if line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('# '):
            # Skip top-level headings (we handle cover page separately)
            pass
        elif line.strip() == '---':
            elements.append(('hr', None))
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            # Bold text as subheading
            text = line.strip()[2:-2]
            elements.append(('bold', text))
        elif line.strip():
            # Regular paragraph
            # Clean up markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
            text = re.sub(r'`([^`]+)`', r'\1', text)  # Code
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
            elements.append(('p', text))

        i += 1

    # Handle any remaining table
    if in_table and table_headers:
        elements.append(('table', (table_headers, table_rows)))

    return elements

def create_document(md_content, output_path):
    """Create the Word document from markdown content."""
    doc = Document()

    # Set up document styles
    setup_styles(doc)

    # Create cover page
    create_cover_page(doc)

    # Create table of contents
    create_toc(doc)

    # Parse and add content
    add_page_break(doc)

    elements = parse_markdown(md_content)

    # Bookmark counter for unique IDs
    bookmark_id = 0

    for elem_type, content in elements:
        if elem_type == 'h2':
            # Major section - add page break before new sections
            if content.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                add_page_break(doc)
            heading = doc.add_heading(content, level=1)

            # Add bookmark for TOC linking
            section_num = content.split('.')[0] if '.' in content else None
            if section_num and section_num.isdigit():
                bookmark_name = f"section_{section_num}"
                # Add bookmark to heading
                run = heading.runs[0] if heading.runs else heading.add_run()
                tag = run._r
                start = OxmlElement('w:bookmarkStart')
                start.set(qn('w:id'), str(bookmark_id))
                start.set(qn('w:name'), bookmark_name)
                tag.insert(0, start)
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), str(bookmark_id))
                tag.append(end)
                bookmark_id += 1

        elif elem_type == 'h3':
            doc.add_heading(content, level=2)
        elif elem_type == 'h4':
            doc.add_heading(content, level=3)
        elif elem_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(11)
        elif elem_type == 'p':
            p = doc.add_paragraph(content)
        elif elem_type == 'code':
            add_code_block(doc, content)
        elif elem_type == 'table':
            headers, rows = content
            add_table(doc, headers, rows)
        elif elem_type == 'hr':
            # Skip horizontal rules - don't add any visual separator
            pass

    # Save document
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

def main():
    # Read markdown file
    input_path = 'ADD_Scarecrow_Drone.md'
    output_path = 'ADD_Scarecrow_Drone_Professional_v2.docx'

    print(f"Reading markdown file: {input_path}")

    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print("Converting to professional DOCX format...")
    create_document(md_content, output_path)

    print("Conversion complete!")
    print(f"\nOutput file: {output_path}")
    print("\nThe document includes:")
    print("  - Professional cover page")
    print("  - Table of contents")
    print("  - Properly formatted tables with borders")
    print("  - ASCII diagrams in monospace font")
    print("  - Consistent heading hierarchy")
    print("  - Page breaks between major sections")

if __name__ == '__main__':
    main()
