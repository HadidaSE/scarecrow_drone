"""
Professional DOCX Generator for ADD_Scarecrow_Drone.md

Creates a professionally formatted Word document with:
- Professional cover page
- Table of contents with clickable links
- Clean tables with borders
- EMBEDDED IMAGES for ASCII diagrams (instead of monospace text)
- Consistent spacing and heading hierarchy
- Skips unimplemented UI mockups (Section 6)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Directory containing generated diagram images
DIAGRAM_DIR = 'diagram_images'

# Mapping of section/subsection names to image files
# This will be built dynamically based on what images exist
IMAGE_MAPPING = {}

def build_image_mapping():
    """Build a mapping of diagram names to image files."""
    global IMAGE_MAPPING

    if not os.path.exists(DIAGRAM_DIR):
        print(f"Warning: {DIAGRAM_DIR} directory not found")
        return

    # Map common names to their image files
    files = os.listdir(DIAGRAM_DIR)

    # Create counters for sequence tracking
    counters = {}

    for f in sorted(files):
        if f.endswith('.png'):
            # Extract base name and number
            base = f.replace('.png', '')
            # Check if it has a number suffix
            match = re.match(r'(.+)_(\d+)$', base)
            if match:
                key = match.group(1)
                num = int(match.group(2))
            else:
                key = base
                num = 1

            if key not in IMAGE_MAPPING:
                IMAGE_MAPPING[key] = []
            IMAGE_MAPPING[key].append((num, os.path.join(DIAGRAM_DIR, f)))

    # Sort each list by number
    for key in IMAGE_MAPPING:
        IMAGE_MAPPING[key].sort(key=lambda x: x[0])

    print(f"Found {sum(len(v) for v in IMAGE_MAPPING.values())} diagram images")

def get_next_image(section_key):
    """Get the next image for a given section key."""
    # Normalize the key
    key = section_key.replace(' ', '_')
    key = re.sub(r'^[\d.]+\s*', '', key)  # Remove leading numbers
    key = re.sub(r'[<>:"/\\|?*]', '', key)
    key = key[:50].strip('_')

    # Direct section name to image file mappings
    explicit_mappings = {
        'Deployment_Diagram': 'Deployment_Diagram',
        'Use_Case_Diagram': 'Use_Case_Diagram',
        'Class_Diagram': 'Class_Diagram',
        'Entity_Relationship_Diagram_ERD': 'Data_Objects_Relationships',
        'Data_Objects_Relationships': 'Data_Objects_Relationships',
        'Database_Schema': 'Databases',
        'Databases': 'Databases',
        'Sequence_Diagrams': 'Sequence_Diagrams',
        'State_Machines': 'State_Machines',
        'Packages': 'Packages',
        'Package_Diagram': 'Packages',
    }

    # Check explicit mapping
    if key in explicit_mappings:
        mapped_key = explicit_mappings[key]
        if mapped_key in IMAGE_MAPPING and IMAGE_MAPPING[mapped_key]:
            _, path = IMAGE_MAPPING[mapped_key].pop(0)
            return path

    if key in IMAGE_MAPPING and IMAGE_MAPPING[key]:
        _, path = IMAGE_MAPPING[key].pop(0)
        return path

    # Try partial matches
    for map_key in list(IMAGE_MAPPING.keys()):
        if key.lower() in map_key.lower() or map_key.lower() in key.lower():
            if IMAGE_MAPPING[map_key]:
                _, path = IMAGE_MAPPING[map_key].pop(0)
                return path

    return None

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def keep_table_on_one_page(table):
    """Keep the entire table on one page by setting keepNext on all rows except the last."""
    for row in table.rows:
        set_row_cant_split(row)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.get_or_add_pPr()
                keepNext = OxmlElement('w:keepNext')
                pPr.append(keepNext)
                keepLines = OxmlElement('w:keepLines')
                pPr.append(keepLines)

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def create_cover_page(doc):
    """Create a professional cover page."""
    for _ in range(8):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Application Design Document (ADD)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scarecrow Drone - Bird Detection System")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version: ")
    run.bold = True
    run.font.size = Pt(12)
    run = version.add_run("1.0")
    run.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Date: ")
    run.bold = True
    run.font.size = Pt(12)
    run = date_p.add_run("January 2026")
    run.font.size = Pt(12)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run("Project: ")
    run.bold = True
    run.font.size = Pt(12)
    run = project.add_run("Scarecrow Drone - Secondary Alpha Implementation")
    run.font.size = Pt(12)

def create_toc(doc):
    """Create table of contents page."""
    add_page_break(doc)

    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    doc.add_paragraph()

    toc_items = [
        "1. Use Cases",
        "2. System Architecture",
        "3. Data Model",
        "4. Behavioral Analysis",
        "5. Object-Oriented Model",
        "6. User Interface Draft",
        "7. Testing"
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.underline = True

def setup_styles(doc):
    """Set up custom styles for the document."""
    styles = doc.styles

    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(10)

    h3 = styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(8)

    h4 = styles['Heading 4']
    h4.font.name = 'Arial'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)

    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_row = table.rows[0]
    header_cells = header_row.cells
    for idx, header in enumerate(headers):
        if idx < len(header_cells):
            header_text = header.strip()
            header_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', header_text)
            header_cells[idx].text = header_text
            for paragraph in header_cells[idx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
            set_cell_shading(header_cells[idx], 'E0E0E0')

    for row_data in rows:
        new_row = table.add_row()
        row_cells = new_row.cells
        for idx, cell_text in enumerate(row_data):
            if idx < len(row_cells):
                cell_text = cell_text.strip()
                cell_text = cell_text.replace('<br>', '\n')
                cell_text = cell_text.replace('&nbsp;', ' ')
                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                row_cells[idx].text = cell_text
                for paragraph in row_cells[idx].paragraphs:
                    paragraph.paragraph_format.space_before = Pt(4)
                    paragraph.paragraph_format.space_after = Pt(4)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'

    # Keep entire table on one page
    keep_table_on_one_page(table)

    doc.add_paragraph()

def add_image(doc, image_path, width_inches=6.5):
    """Add an image to the document, centered."""
    if not os.path.exists(image_path):
        print(f"  Warning: Image not found: {image_path}")
        return False

    # Add paragraph for the image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # Add space after image
    doc.add_paragraph()
    return True

def add_code_block(doc, code_text, smaller_font=False):
    """Add a code block (for OCL code or file structure) to the document."""
    lines = code_text.split('\n')
    # Use 7pt for file structures to fit on one page, 9pt for OCL code
    font_size = 7 if smaller_font else 9

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

def is_ascii_diagram(code_text):
    """Check if code block is an ASCII diagram (vs OCL or other code)."""
    # OCL indicators
    ocl_keywords = ['context', 'pre:', 'post:', 'inv:', 'self.', 'implies', 'forAll', 'exists']
    text_lower = code_text.lower()

    for keyword in ocl_keywords:
        if keyword.lower() in text_lower:
            return False

    # Check for box-drawing characters
    box_chars = '┌┐└┘│─┬┴├┤╔╗╚╝║═+|'
    has_box = any(c in code_text for c in box_chars)

    return has_box

def is_file_structure(code_text):
    """Check if code block is a file structure tree."""
    # File structures have specific patterns like:
    # ├── folder/
    # │   ├── file.py
    # └── another_file.js

    # Must have tree-drawing characters followed by a space and a name
    # (not just decorative box corners like └────────────────┘)
    import re

    # Look for tree pattern followed by a word (filename/folder)
    # Pattern: tree char + dashes + space + word
    tree_file_pattern = r'[├└]── \w'

    # Also need multiple levels (│   ├ or similar)
    has_nested = '│   ' in code_text or '│  ├' in code_text or '│  └' in code_text

    has_tree_file = bool(re.search(tree_file_pattern, code_text))

    return has_tree_file and has_nested

def is_ui_mockup(code_text):
    """Check if code block is a UI mockup (ASCII art interface)."""
    # UI mockups have lots of + and - for borders
    plus_count = code_text.count('+')
    dash_count = code_text.count('-')
    pipe_count = code_text.count('|')

    # If it looks like a UI mockup (many borders, boxes)
    if plus_count > 10 and (dash_count > 50 or pipe_count > 30):
        return True

    # Also check for common UI elements
    ui_elements = ['[Dashboard]', '[Help]', '[Back]', '[Save', '[Cancel]',
                   '[Start', '[Stop', '[Apply', '[View', 'Interface:',
                   'Purpose:', 'Inputs:', 'Outputs:']
    return any(elem in code_text for elem in ui_elements)

def parse_markdown(md_content):
    """Parse markdown content and convert to structured data."""
    lines = md_content.split('\n')
    elements = []

    i = 0
    in_code_block = False
    code_content = []
    code_type = None
    in_table = False
    table_headers = []
    table_rows = []
    skip_google_docs_format = False
    current_h3 = ""
    current_h2 = ""
    skip_section_6_content = False

    while i < len(lines):
        line = lines[i]

        # Track section 6 (User Interface Draft) to skip entirely
        # These interfaces are not implemented yet
        if line.startswith('## 6.'):
            skip_section_6_content = True
            i += 1
            continue
        elif line.startswith('## 7.'):
            skip_section_6_content = False

        # Skip all content in Section 6
        if skip_section_6_content:
            i += 1
            continue

        # Skip Google Docs format sections
        if '**Google Docs Format:**' in line or 'Google Docs Format:' in line:
            skip_google_docs_format = True
            i += 1
            continue

        if skip_google_docs_format:
            if line.startswith('##') or line.startswith('###') or line.strip() == '---':
                skip_google_docs_format = False
            else:
                i += 1
                continue

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    code_text = '\n'.join(code_content)

                    if is_file_structure(code_text):
                        # File structure - render as smaller code
                        elements.append(('file_structure', code_text))
                    elif is_ascii_diagram(code_text):
                        # It's an ASCII diagram - we'll use an image
                        elements.append(('diagram', (current_h3, code_text)))
                    else:
                        # It's OCL or other code - keep as text
                        elements.append(('code', code_text))
                code_content = []
                in_code_block = False
            else:
                # Start code block - check if it's OCL
                code_type = line.strip()[3:].strip()  # Get language after ```
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and not line.strip().startswith('#'):
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]

            if not in_table:
                in_table = True
                table_headers = cells
            elif re.match(r'^[\s|:\-]+$', line.replace('|', '').strip()):
                pass
            else:
                table_rows.append(cells)

            i += 1
            continue
        elif in_table:
            if table_headers:
                elements.append(('table', (table_headers, table_rows)))
            table_headers = []
            table_rows = []
            in_table = False

        # Handle headings
        if line.startswith('#### '):
            # Skip h4 headings - they are often redundant
            pass
        elif line.startswith('### '):
            current_h3 = line[4:].strip()
            elements.append(('h3', current_h3))
        elif line.startswith('## '):
            current_h2 = line[3:].strip()
            elements.append(('h2', current_h2))
        elif line.startswith('# '):
            pass
        elif line.strip() == '---':
            elements.append(('hr', None))
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            elements.append(('bold', text))
        elif line.strip():
            text = line.strip()
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            elements.append(('p', text))

        i += 1

    if in_table and table_headers:
        elements.append(('table', (table_headers, table_rows)))

    return elements

def create_document(md_content, output_path):
    """Create the Word document from markdown content."""
    doc = Document()

    setup_styles(doc)
    create_cover_page(doc)
    create_toc(doc)
    add_page_break(doc)

    elements = parse_markdown(md_content)

    # Track image usage per section
    image_counters = {}
    diagram_count = 0
    image_count = 0

    for elem_type, content in elements:
        if elem_type == 'h2':
            if content.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                add_page_break(doc)
            doc.add_heading(content, level=1)
        elif elem_type == 'h3':
            doc.add_heading(content, level=2)
        elif elem_type == 'h4':
            doc.add_heading(content, level=3)
        elif elem_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(11)
        elif elem_type == 'p':
            p = doc.add_paragraph(content)
        elif elem_type == 'code':
            add_code_block(doc, content)
        elif elem_type == 'file_structure':
            # File structure with smaller font to fit on one page
            add_code_block(doc, content, smaller_font=True)
        elif elem_type == 'diagram':
            section_name, ascii_content = content
            diagram_count += 1

            # Debug print
            key = section_name.replace(' ', '_')
            key = re.sub(r'^[\d.]+\s*', '', key)
            key = key.strip('_')
            print(f"  Diagram section: {section_name!r} -> key: {key!r}")

            # Try to find matching image
            # Special case: Deployment Diagram was split into 2 images - add both
            image_path = get_next_image(section_name)
            images_added = 0

            if image_path and os.path.exists(image_path):
                if add_image(doc, image_path, width_inches=6.5):
                    image_count += 1
                    images_added += 1
                    print(f"  Added image: {image_path}")

                # For Deployment Diagram only, also add the second part
                if 'Deployment' in section_name:
                    image_path2 = get_next_image(section_name)
                    if image_path2 and os.path.exists(image_path2):
                        if add_image(doc, image_path2, width_inches=6.5):
                            image_count += 1
                            images_added += 1
                            print(f"  Added image: {image_path2}")

            if images_added == 0:
                # No images found - use ASCII as fallback
                print(f"  No image for: {section_name}, using ASCII")
                add_code_block(doc, ascii_content)

        elif elem_type == 'table':
            headers, rows = content
            add_table(doc, headers, rows)
        elif elem_type == 'hr':
            pass

    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print(f"Diagrams found: {diagram_count}")
    print(f"Images embedded: {image_count}")

def main():
    input_path = 'ADD_Scarecrow_Drone.md'
    output_path = 'ADD_Scarecrow_Drone_Professional.docx'

    print(f"Reading markdown file: {input_path}")

    # Build image mapping
    build_image_mapping()

    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print("Converting to professional DOCX format with images...")
    create_document(md_content, output_path)

    print("\nConversion complete!")
    print(f"Output file: {output_path}")
    print("\nThe document includes:")
    print("  - Professional cover page")
    print("  - Table of contents")
    print("  - Properly formatted tables with borders")
    print("  - Embedded diagram images (replacing ASCII art)")
    print("  - OCL code blocks in monospace font")
    print("  - Consistent heading hierarchy")
    print("  - Page breaks between major sections")
    print("  - UI mockups removed (Section 6)")

if __name__ == '__main__':
    main()
