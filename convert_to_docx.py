import re
from docx import Document
from docx.shared import Pt

# Read markdown file
with open('ADD_Scarecrow_Drone.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Create Word document
doc = Document()

# Process markdown line by line
lines = md_content.split('\n')
i = 0
in_code_block = False
code_content = []
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i]

    # Handle code blocks
    if line.startswith('```'):
        if in_code_block:
            # End code block - add as paragraph with monospace
            code_text = '\n'.join(code_content)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            code_content = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_content.append(line)
        i += 1
        continue

    # Handle tables
    if '|' in line and not line.strip().startswith('#'):
        if not in_table:
            in_table = True
            table_rows = []
        # Skip separator lines
        if re.match(r'^[\s|:\-]+$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            table_rows.append(cells)
        i += 1
        continue
    elif in_table:
        # End of table
        if table_rows:
            num_cols = max(len(row) for row in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = 'Table Grid'
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < num_cols:
                        table.rows[row_idx].cells[col_idx].text = cell_data
            doc.add_paragraph()
        table_rows = []
        in_table = False

    # Handle headers
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=3)
    elif line.startswith('---'):
        doc.add_paragraph('_' * 50)
    elif line.strip():
        doc.add_paragraph(line)

    i += 1

# Handle any remaining table
if in_table and table_rows:
    num_cols = max(len(row) for row in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = 'Table Grid'
    for row_idx, row_data in enumerate(table_rows):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < num_cols:
                table.rows[row_idx].cells[col_idx].text = cell_data

doc.save('ADD_Scarecrow_Drone.docx')
print('Created ADD_Scarecrow_Drone.docx')
